"""ingest-001 -- PDF in, page-anchored chunks out.

The whole citation story rests on one invariant, enforced here and nowhere else:

    normalise(page_text(pdf, chunk.page_no))[chunk.char_start:chunk.char_end]
        == chunk.text

Every chunk is a *slice* of exactly one page's normalised text. It is never
rebuilt, re-joined or re-flowed after the span is taken, and it never spans two
pages -- a chunk straddling a page boundary could not honestly claim either
page number. `verify_material()` re-derives that invariant straight from the
PDF, which is what turns "Show Source" from a claim into a check.

No model is ever asked where a passage came from. That is why our citations
cannot hallucinate.

Split in two on purpose:

  * `parse_pdf()` and everything above it is pure -- no database, no
    embeddings, no network. That is the half the test suite exercises, against
    a PDF it builds itself.
  * `ingest_material()` downwards is the half that touches Neon and fastembed.
"""

from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

# Chunk geometry. Tuned for bge-small's 512-token window: ~800 characters is
# roughly 180 tokens of prose, comfortably inside it, and long enough that a
# retrieved passage reads as an explanation rather than a fragment.
TARGET_CHARS = 800
MAX_CHARS = 1400
OVERLAP_SENTENCES = 1
MIN_CHUNK_CHARS = 80
MIN_PAGE_CHARS = 40      # below this a page is a cover, a plate, or blank

SOFT_HYPHEN = "­"


# ---------------------------------------------------------------------------
# Text normalisation
# ---------------------------------------------------------------------------

def normalise(raw: str) -> str:
    """PDF text layer -> flat prose, deterministically.

    Must stay a pure function of its input: `verify_material()` re-runs it long
    after ingestion and compares the result against what was stored, so any
    dependence on time, locale or configuration would break the audit.

    Paragraph breaks survive as a blank line; soft wraps inside a paragraph
    become spaces, because a line break mid-sentence is a typesetting artefact
    and storing it would put "force.\\nThe" inside a passage a student reads.

    NFKC first, because typeset books are full of characters that are not the
    letters they look like: "eﬃcient" is a single ligature glyph, so a student
    searching for "efficient" would miss the page it is on, and the snippet
    would render as mojibake outside a Unicode-clean terminal. NFKC also folds
    non-breaking spaces into ordinary ones, which is what stops "array p" from
    arriving as one unsplittable token.
    """
    t = unicodedata.normalize("NFKC", raw)
    t = t.replace("\r\n", "\n").replace("\r", "\n")
    t = t.replace(SOFT_HYPHEN, "")
    t = re.sub(r"-\n(?=[a-z])", "", t)        # word hyphenated across a line break
    t = re.sub(r"\n{2,}", "\x00", t)          # protect paragraph breaks
    t = t.replace("\n", " ")                  # soft wraps become spaces
    t = t.replace("\x00", "\n\n")
    t = re.sub(r"[^\S\n]+", " ", t)           # runs of spaces/tabs/NBSP -> one space
    t = re.sub(r" *\n\n *", "\n\n", t)
    return t.strip()


def squash(text: str) -> str:
    """Whitespace-free lowercase form, for comparisons that must not care about
    layout.

    Verification uses it to match a stored chunk against a FRESH extraction of
    the page, independently of the char span we recorded -- so the check cannot
    pass merely because ingestion was self-consistently wrong.

    It folds the same way `normalise()` does, and for the same reason: one side
    of that comparison is raw PDF text where "ﬃ" is still one glyph, and a
    mismatch there would be a Unicode artefact, not a citation error.
    """
    folded = unicodedata.normalize("NFKC", text).replace(SOFT_HYPHEN, "")
    folded = folded.replace("\r\n", "\n").replace("\r", "\n")
    # The line-break hyphen has to go BEFORE whitespace is stripped, and it has
    # to go for the same reason `normalise` drops it: "pre-\nceding" is one
    # word that a typesetter broke, not a hyphenated compound.
    #
    # Without this, squash() folded one side of the comparison differently from
    # the other -- the stored chunk came from normalised text ("preceding")
    # while verification squashed the RAW page ("pre-ceding") -- and every
    # chunk containing a line-broken word failed verification even though its
    # page number was correct. Three chunks were reported as citation errors
    # that were not, one of them in a textbook.
    folded = re.sub(r"-\n(?=[a-z])", "", folded)
    return re.sub(r"\s+", "", folded).lower()


# ---------------------------------------------------------------------------
# Sentence-aligned chunking, by slicing only
# ---------------------------------------------------------------------------

_BOUNDARY = re.compile(r"(?<=[.!?:;])\s+|\n\n")


def sentence_spans(text: str) -> list[tuple[int, int]]:
    """(start, end) of each sentence, in order. Any run of them concatenates
    back into a single contiguous slice of `text` -- that property is what lets
    a chunk be a slice rather than a rebuild."""
    spans: list[tuple[int, int]] = []
    pos = 0
    for m in _BOUNDARY.finditer(text):
        if m.start() > pos:
            spans.append((pos, m.start()))
        pos = m.end()
    if pos < len(text):
        spans.append((pos, len(text)))
    return spans


def _hard_split(text: str, start: int, end: int) -> list[tuple[int, int]]:
    """A 'sentence' longer than MAX_CHARS: a table, a formula dump, a page with
    no punctuation at all. Cut at spaces so we are still only ever slicing."""
    out: list[tuple[int, int]] = []
    while end - start > MAX_CHARS:
        cut = text.rfind(" ", start + MIN_CHUNK_CHARS, start + MAX_CHARS)
        if cut == -1:
            cut = start + MAX_CHARS
        out.append((start, cut))
        start = cut + 1
    if end > start:
        out.append((start, end))
    return out


def chunk_spans(text: str) -> list[tuple[int, int]]:
    """Group sentences into ~TARGET_CHARS chunks with one sentence of overlap.

    The overlap exists so a definition that lands at the end of one chunk is
    still retrievable alongside the sentence that uses it.
    """
    units: list[tuple[int, int]] = []
    for s, e in sentence_spans(text):
        units.extend(_hard_split(text, s, e) if e - s > MAX_CHARS else [(s, e)])
    if not units:
        return []

    chunks: list[tuple[int, int]] = []
    i = 0
    while i < len(units):
        start = units[i][0]
        j = i
        while j < len(units) - 1 and units[j][1] - start < TARGET_CHARS:
            j += 1
        chunks.append((start, units[j][1]))
        if j >= len(units) - 1:
            break
        i = max(j + 1 - OVERLAP_SENTENCES, i + 1)

    # A stub tail reads as noise and pollutes retrieval; fold it into its
    # predecessor. Still a slice: the spans are contiguous.
    if len(chunks) > 1 and chunks[-1][1] - chunks[-1][0] < MIN_CHUNK_CHARS:
        last = chunks.pop()
        chunks[-1] = (chunks[-1][0], last[1])
    return chunks


# ---------------------------------------------------------------------------
# PDF -> drafts
# ---------------------------------------------------------------------------

@dataclass(frozen=True)
class ChunkDraft:
    """One future Chunk row. page_no is 1-based -- what a human reads off the
    page, and what "Show Source" prints."""

    page_no: int
    chapter: str | None
    char_start: int
    char_end: int
    text: str


@dataclass
class ParsedDoc:
    page_count: int
    pages_with_text: int
    chunks: list[ChunkDraft]

    @property
    def chunk_count(self) -> int:
        return len(self.chunks)


def _open(path: str | Path):
    """Imported lazily so importing `app.main` (and the test suite) does not
    pay for PyMuPDF, and does not need it installed at all."""
    import pymupdf   # `fitz` is the deprecated spelling of the same module

    return pymupdf.open(str(path))


def chapter_index(doc) -> list[tuple[int, str]]:
    """Top-level table of contents as (first_page, title), page-ordered.

    Read from the PDF's own outline. A book without an outline yields nothing
    and `chapter` stays NULL -- a guessed chapter heading is worse than an
    absent one, because it shows up on a citation looking like fact.
    """
    try:
        toc = doc.get_toc(simple=True) or []
    except Exception:                     # noqa: BLE001 - a broken outline must
        return []                         # not fail an otherwise fine ingest
    out = [
        (int(page), str(title).strip())
        for level, title, page in toc
        if level == 1 and int(page) >= 1 and str(title).strip()
    ]
    return sorted(out)


def chapter_for_page(index: list[tuple[int, str]], page_no: int) -> str | None:
    found = None
    for start, title in index:
        if start <= page_no:
            found = title
        else:
            break
    return found


def page_text(path: str | Path, page_no: int) -> str:
    """Normalised text of one 1-based page. The verification entry point."""
    doc = _open(path)
    try:
        return normalise(doc[page_no - 1].get_text("text"))
    finally:
        doc.close()


def parse_pdf(path: str | Path) -> ParsedDoc:
    """Pure: no database, no embeddings, no network."""
    doc = _open(path)
    try:
        index = chapter_index(doc)
        chunks: list[ChunkDraft] = []
        pages_with_text = 0

        for i in range(doc.page_count):
            page_no = i + 1
            text = normalise(doc[i].get_text("text"))
            if len(text) < MIN_PAGE_CHARS:
                continue
            pages_with_text += 1
            chapter = chapter_for_page(index, page_no)
            for start, end in chunk_spans(text):
                chunks.append(ChunkDraft(
                    page_no=page_no,
                    chapter=chapter,
                    char_start=start,
                    char_end=end,
                    text=text[start:end],
                ))
        return ParsedDoc(doc.page_count, pages_with_text, chunks)
    finally:
        doc.close()


# ---------------------------------------------------------------------------
# Drafts -> database  (this half needs Neon and fastembed)
# ---------------------------------------------------------------------------

EMBED_BATCH = 64
# Derived, not restated. This list lived in four places -- models, the admin
# router, this module and the CLI -- and admin-007 found them disagreeing:
# adding "reference" to the model left the CLI rejecting it. Imported here
# rather than at the top of the file to keep the parsing half above genuinely
# free of app imports, as the module docstring promises.
from ..models import MATERIAL_KINDS  # noqa: E402

INGESTABLE_KINDS = MATERIAL_KINDS


@dataclass
class IngestResult:
    material_id: int
    title: str
    kind: str
    page_count: int
    pages_with_text: int
    chunk_count: int
    replaced: int          # chunks deleted before re-ingesting this material


def ingest_material(
    db,
    *,
    course_id: int,
    title: str,
    kind: str,
    path: str | Path,
    uploaded_by_id: int | None = None,
    progress=None,
) -> IngestResult:
    """Parse, chunk, embed and store one PDF. Idempotent on (course_id, title).

    Re-ingesting replaces the material's chunks rather than appending: running
    it twice must not double every passage in the corpus and quietly bias
    retrieval towards whichever book was ingested most often.

    **This function commits, and it cannot be composed inside a caller's
    transaction.** Embedding a whole book takes minutes, and a hosted Postgres
    terminates a session that sits idle inside an open transaction -- Neon did
    exactly that on the first real run, after all 1240 inserts, rolling back the
    entire book on the final UPDATE. So the order here is deliberate:

        1. parse and embed with NO transaction open,
        2. then write, committing each batch.

    Any uncommitted work the caller was holding is committed at step 1, because
    leaving it open is what causes the failure.

    `ingest_status` is 'running' from the first write until the last, so a
    process killed midway leaves a row that says so rather than one that claims
    to be complete.
    """
    from sqlalchemy import delete, select

    from ..models import MATERIAL_KINDS, Chunk, Material
    from .embed import embed_documents

    if kind not in INGESTABLE_KINDS:
        raise ValueError(f"kind must be one of {INGESTABLE_KINDS}, got {kind!r}")

    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(path)

    parsed = parse_pdf(path)
    if not parsed.chunks:
        raise ValueError(
            f"{path.name} produced no text. It is probably a scanned image with "
            f"no text layer -- OCR is out of scope for this build."
        )

    def report(phase: str, done: int) -> None:
        if progress:
            progress(phase, done, parsed.chunk_count)

    # --- 1. the slow part, with nothing held ------------------------------
    db.commit()
    vectors: list[list[float]] = []
    for start in range(0, len(parsed.chunks), EMBED_BATCH):
        batch = parsed.chunks[start:start + EMBED_BATCH]
        vectors.extend(embed_documents([c.text for c in batch]))
        report("embedding", len(vectors))

    # --- 2. the writes, in committed batches ------------------------------
    material = db.scalar(
        select(Material).where(Material.course_id == course_id, Material.title == title)
    )
    if material is None:
        material = Material(course_id=course_id, title=title)
        db.add(material)
    material.kind = kind
    material.page_count = parsed.page_count
    material.source_path = path.name       # name only: absolute paths differ per machine
    material.status = "active"
    material.ingest_status = "running"
    material.chunk_count = 0
    if uploaded_by_id is not None:
        material.uploaded_by_id = uploaded_by_id
    db.flush()

    replaced = db.query(Chunk).filter(Chunk.material_id == material.id).count()
    db.execute(delete(Chunk).where(Chunk.material_id == material.id))
    db.commit()

    for start in range(0, len(parsed.chunks), EMBED_BATCH):
        for draft, vec in zip(parsed.chunks[start:start + EMBED_BATCH],
                              vectors[start:start + EMBED_BATCH]):
            db.add(Chunk(
                material_id=material.id,
                page_no=draft.page_no,
                chapter=draft.chapter,
                char_start=draft.char_start,
                char_end=draft.char_end,
                text=draft.text,
                embedding=vec,
            ))
        db.commit()
        report("writing", min(start + EMBED_BATCH, parsed.chunk_count))

    material.chunk_count = parsed.chunk_count
    material.ingest_status = "complete"
    db.commit()

    return IngestResult(
        material_id=material.id,
        title=title,
        kind=kind,
        page_count=parsed.page_count,
        pages_with_text=parsed.pages_with_text,
        chunk_count=parsed.chunk_count,
        replaced=replaced,
    )


# ---------------------------------------------------------------------------
# Verification -- ingest-001's third acceptance step, as code
# ---------------------------------------------------------------------------

@dataclass
class ChunkCheck:
    chunk_id: int
    page_no: int
    span: tuple[int, int]
    slice_matches: bool      # stored text == that span of the re-extracted page
    appears_on_page: bool    # stored text occurs on the page, ignoring layout
    preview: str

    @property
    def ok(self) -> bool:
        return self.slice_matches and self.appears_on_page


def verify_material(db, material_id: int, pdf_path: str | Path,
                    sample: int = 3, seed: int = 0) -> list[ChunkCheck]:
    """Re-open the PDF and prove sampled chunks really came off the page they claim.

    Two independent checks per chunk:

      1. the recorded span, re-sliced out of a fresh extraction, is byte-identical
         to the stored text;
      2. the stored text occurs in that page's raw extraction with all whitespace
         removed -- which does not use our span at all, so a wrong char_start
         cannot slip through by agreeing with itself.

    `seed` is fixed by default so the evidence file is reproducible.

    Sampling happens in two cheap queries -- ids, then the three rows we want --
    rather than by loading the material and picking in Python. A book is over a
    thousand chunks and every row carries a 384-float embedding, so pulling them
    all to keep three is megabytes over the wire; it took long enough that Neon
    closed the connection mid-verification. The columns fetched here are only
    the ones the check reads, so the embeddings never leave the database.
    """
    import random

    from sqlalchemy import select

    from ..models import Chunk

    ids = db.scalars(
        select(Chunk.id).where(Chunk.material_id == material_id).order_by(Chunk.id)
    ).all()
    if not ids:
        return []

    picked = random.Random(seed).sample(list(ids), min(sample, len(ids)))
    rows = db.execute(
        select(Chunk.id, Chunk.page_no, Chunk.char_start, Chunk.char_end, Chunk.text)
        .where(Chunk.id.in_(picked))
        .order_by(Chunk.page_no, Chunk.char_start)
    ).all()

    doc = _open(pdf_path)
    try:
        checks = []
        for chunk_id, page_no, char_start, char_end, text in rows:
            raw = doc[page_no - 1].get_text("text")
            page = normalise(raw)
            checks.append(ChunkCheck(
                chunk_id=chunk_id,
                page_no=page_no,
                span=(char_start, char_end),
                slice_matches=page[char_start:char_end] == text,
                appears_on_page=squash(text) in squash(raw),
                preview=" ".join(text.split())[:90],
            ))
        return checks
    finally:
        doc.close()


def stub_corpus_materials(db, course_id: int | None = None) -> list[str]:
    """Titles still coming from the stand-in corpus.json rather than a PDF.

    Placeholder passages compete with the real book in search results, so
    ingestion warns when both are present *in the same course*. Retrieval is
    course-scoped, so a stub in another course is not competing with anything --
    passing `course_id` keeps the warning actionable rather than noise.
    """
    from sqlalchemy import select

    from ..models import Course, Material

    stmt = (
        select(Material.title, Course.code)
        .join(Course, Course.id == Material.course_id)
        .where(Material.source_path.is_(None), Material.ingest_status == "complete")
    )
    if course_id is not None:
        stmt = stmt.where(Material.course_id == course_id)
    return [f"{title}  ({code})" for title, code in db.execute(stmt).all()]


# ---------------------------------------------------------------------------
# Reflowable sources (EPUB and friends) -> a fixed-page PDF
# ---------------------------------------------------------------------------

REFLOWABLE_SUFFIXES = {".epub", ".mobi", ".fb2"}
# admin-007. PyMuPDF opens the set above directly; these three it cannot, so
# they are read to text first and laid out the same way. They are reflowable in
# exactly the same sense -- no page exists until we fix a paper size.
TEXT_SUFFIXES = {".txt", ".md"}
DOCX_SUFFIXES = {".docx"}
LAYOUT_PAPER = "a4"
LAYOUT_FONTSIZE = 11


def supported_suffixes() -> set[str]:
    """Everything ingestion can take. One place, so the upload endpoint, the
    folder scan and the error message cannot disagree about it."""
    return {".pdf", *REFLOWABLE_SUFFIXES, *TEXT_SUFFIXES, *DOCX_SUFFIXES}


def _read_text_source(path: Path) -> str:
    """Plain text out of a .txt/.md/.docx.

    A .docx is a zip of XML, not something PyMuPDF can open -- python-docx
    pulls the paragraphs out. Tables are included because an assignment's
    questions are frequently in one, and dropping them would silently ingest
    half a document.
    """
    if path.suffix.lower() in DOCX_SUFFIXES:
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    # errors="replace": a stray byte in a teacher's notes must not abort an
    # ingest that is otherwise fine.
    return path.read_text(encoding="utf-8", errors="replace")


def _text_to_fixed_pdf(text: str, target: Path) -> None:
    """Lay text out at A4/11pt so it has pages to cite.

    Same bargain as an EPUB: the page numbers are reproducible and they are
    OURS, not a publisher's.
    """
    import pymupdf

    story = pymupdf.Story(html=_text_as_html(text))
    writer = pymupdf.DocumentWriter(str(target))
    rect = pymupdf.paper_rect(LAYOUT_PAPER) + (36, 36, -36, -36)  # 0.5in margins
    more = True
    while more:
        device = writer.begin_page(pymupdf.paper_rect(LAYOUT_PAPER))
        more, _ = story.place(rect)
        story.draw(device)
        writer.end_page()
    writer.close()


def _text_as_html(text: str) -> str:
    """Paragraphs, escaped. No markdown rendering on purpose -- a heading that
    became an <h1> would change the page count depending on how the source was
    marked up, and page numbers have to be stable."""
    from html import escape

    paragraphs = [escape(block.strip()).replace("\n", "<br/>")
                  for block in text.split("\n\n") if block.strip()]
    body = "".join(f"<p>{p}</p>" for p in paragraphs) or "<p></p>"
    return (f"<html><body style='font-family:sans-serif;"
            f"font-size:{LAYOUT_FONTSIZE}pt'>{body}</body></html>")


def ensure_fixed_pdf(path: str | Path, out_dir: str | Path) -> tuple[Path, bool]:
    """A path we can cite page numbers from. Returns (pdf_path, converted).

    An EPUB has no pages -- it has a stream of text that a reader paginates on
    the fly, so "page 143" means nothing until someone fixes a paper size and a
    font size. We fix them here (A4 / 11pt), convert once, and cite the result.

    That makes our page numbers *reproducible* but not the publisher's. A
    citation from a converted EPUB points at page 143 of this A4 rendering, not
    page 143 of the printed book. Where a real PDF of the same title exists,
    prefer it -- its page numbers are the ones a student can check against a
    physical copy.

    Conversion is skipped when the .pdf is already there and newer than the
    source, so re-running ingestion is cheap.
    """
    import pymupdf

    path, out_dir = Path(path), Path(out_dir)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return path, False
    if suffix not in supported_suffixes():
        raise ValueError(
            f"Cannot ingest {path.suffix!r}. Supported: "
            f"{', '.join(sorted(supported_suffixes()))}"
        )

    out_dir.mkdir(parents=True, exist_ok=True)
    target = out_dir / f"{path.stem}.pdf"
    if target.exists() and target.stat().st_mtime >= path.stat().st_mtime:
        return target, False

    # admin-007: text-ish sources PyMuPDF cannot open. Read to text, then lay
    # out exactly as an EPUB is, so page attribution works the same way.
    if suffix in TEXT_SUFFIXES or suffix in DOCX_SUFFIXES:
        _text_to_fixed_pdf(_read_text_source(path), target)
        return target, True

    src = pymupdf.open(str(path))
    try:
        src.layout(rect=pymupdf.paper_rect(LAYOUT_PAPER), fontsize=LAYOUT_FONTSIZE)
        toc = src.get_toc(simple=True)
        fixed = pymupdf.open("pdf", src.convert_to_pdf())
        try:
            # convert_to_pdf() drops the outline, and the outline is where our
            # chapter names come from -- carry it across explicitly.
            if toc:
                fixed.set_toc(toc)
            fixed.save(str(target), garbage=3, deflate=True)
        finally:
            fixed.close()
    finally:
        src.close()
    return target, True
